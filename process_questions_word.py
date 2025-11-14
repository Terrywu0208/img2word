import re
import os
import csv
import zipfile
from PIL import Image
import pytesseract
import io
import xml.etree.ElementTree as ET

# Configure Tesseract path for Windows
def configure_tesseract():
    """Configure Tesseract path if not in PATH"""
    try:
        # Try to get Tesseract version to check if it's in PATH
        pytesseract.get_tesseract_version()
    except:
        # Tesseract not found in PATH, try common Windows installation paths
        username = os.getenv('USERNAME', '')
        common_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        if username:
            common_paths.append(rf'C:\Users\{username}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe')
        
        for path in common_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.tesseract_cmd = path
                print(f"Found Tesseract at: {path}")
                return True
        
        # If still not found, provide instructions
        print("ERROR: Tesseract OCR is not installed or not in PATH.")
        print("Please install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")
        print("Or set the path manually by setting: pytesseract.pytesseract.tesseract_cmd = 'path/to/tesseract.exe'")
        return False
    
    return True

# Configure Tesseract on import
configure_tesseract()
try:
    import win32com.client
    WIN32COM_AVAILABLE = True
except ImportError:
    WIN32COM_AVAILABLE = False
    print("Warning: win32com is not installed. Please run: pip install pywin32")
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx is not installed. Please run: pip install python-docx")

def extract_images_from_docx(word_file_path):
    """
    Extract all images from Word document (.docx), in the order they appear in the document
    Returns: (document text content, image list)
    Image list format: [(image index, image binary data, image format), ...]
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required to process .docx files. Please run: pip install python-docx")
    
    # Read text content (including paragraphs and tables)
    doc = Document(word_file_path)
    doc_text_parts = []
    
    # Extract paragraph text
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            doc_text_parts.append(paragraph.text)
    
    # Extract table text (if questions are in tables)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                doc_text_parts.append('\t'.join(row_texts))
    
    doc_text = '\n'.join(doc_text_parts)
    
    images = []
    image_index = 0
    
    # Extract images from ZIP file in the order they appear in the document
    with zipfile.ZipFile(word_file_path, 'r') as zip_ref:
        # Read document.xml to find the order of image references
        try:
            document_xml = zip_ref.read('word/document.xml')
            root = ET.fromstring(document_xml)
            
            # Read relationships file to establish mapping from rId to file path
            relationships = {}
            try:
                rels_xml = zip_ref.read('word/_rels/document.xml.rels')
                rels_root = ET.fromstring(rels_xml)
                for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                    rid = rel.get('Id')
                    target = rel.get('Target')
                    if target and target.startswith('media/'):
                        relationships[rid] = 'word/' + target
            except:
                pass
            
            # Find all image references in document order
            image_rids = []
            # Find all blip elements (image references)
            # Use full namespace URL for search
            blip_ns = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            embed_attr = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
            for blip in root.findall(f'.//{blip_ns}'):
                rid = blip.get(embed_attr)
                if rid and rid not in image_rids:
                    image_rids.append(rid)
            
            # If image references found through XML parsing, extract in order
            if image_rids:
                for rid in image_rids:
                    if rid in relationships:
                        media_file = relationships[rid]
                        if media_file in zip_ref.namelist():
                            image_index += 1
                            try:
                                # Extract file extension
                                _, ext = os.path.splitext(media_file)
                                ext = ext.lstrip('.').lower()
                                
                                # Only process image formats
                                if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
                                    # Read image data
                                    img_data = zip_ref.read(media_file)
                                    # Convert to PNG format uniformly
                                    if ext != 'png':
                                        # Convert to PNG format
                                        img = Image.open(io.BytesIO(img_data))
                                        png_bytes = io.BytesIO()
                                        img.save(png_bytes, format='PNG')
                                        img_data = png_bytes.getvalue()
                                        ext = 'png'
                                    
                                    images.append((image_index, img_data, ext))
                            except Exception as e:
                                print(f"Warning: Error reading image {media_file}: {e}")
            
            # If XML parsing didn't find images, fall back to extracting by filename order (keep original logic as backup)
            if len(images) == 0:
                file_list = zip_ref.namelist()
                media_files = [f for f in file_list if f.startswith('word/media/')]
                # Don't sort, keep original order in ZIP file (usually insertion order)
                for media_file in media_files:
                    _, ext = os.path.splitext(media_file)
                    ext = ext.lstrip('.').lower()
                    if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
                        image_index += 1
                        try:
                            img_data = zip_ref.read(media_file)
                            if ext != 'png':
                                img = Image.open(io.BytesIO(img_data))
                                png_bytes = io.BytesIO()
                                img.save(png_bytes, format='PNG')
                                img_data = png_bytes.getvalue()
                                ext = 'png'
                            images.append((image_index, img_data, ext))
                        except Exception as e:
                            print(f"Warning: Error reading image {media_file}: {e}")
        
        except Exception as e:
            # If XML parsing fails, fall back to extracting by filename order
            print(f"Warning: Failed to parse document XML, will extract images by filename order: {e}")
            file_list = zip_ref.namelist()
            media_files = [f for f in file_list if f.startswith('word/media/')]
            # Don't sort, keep original order in ZIP file
            for media_file in media_files:
                _, ext = os.path.splitext(media_file)
                ext = ext.lstrip('.').lower()
                if ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp']:
                    image_index += 1
                    try:
                        img_data = zip_ref.read(media_file)
                        if ext != 'png':
                            img = Image.open(io.BytesIO(img_data))
                            png_bytes = io.BytesIO()
                            img.save(png_bytes, format='PNG')
                            img_data = png_bytes.getvalue()
                            ext = 'png'
                        images.append((image_index, img_data, ext))
                    except Exception as e2:
                        print(f"Warning: Error reading image {media_file}: {e2}")
    
    return doc_text, images

def extract_images_from_doc(word_file_path):
    """
    Extract all images from Word document (.doc) (legacy format)
    Returns: (document text content, image list)
    Image list format: [(image index, image binary data, image format), ...]
    """
    if not WIN32COM_AVAILABLE:
        raise ImportError("pywin32 is required to process .doc files. Please run: pip install pywin32")
    
    # Create Word application object
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False
    
    try:
        # Open Word document
        doc_path = os.path.abspath(word_file_path)
        doc = word_app.Documents.Open(doc_path)
        
        # Get document text content
        doc_text = doc.Content.Text
        
        # Extract all images
        images = []
        image_index = 0
        temp_dir = os.path.join(os.path.dirname(doc_path), 'temp_images')
        if not os.path.exists(temp_dir):
            os.makedirs(temp_dir)
        
        # Iterate through all InlineShape objects (images) in the document
        for shape_idx in range(1, doc.InlineShapes.Count + 1):
            inline_shape = doc.InlineShapes(shape_idx)
            
            # Check if it's an image (Type = 3 means image)
            if inline_shape.Type == 3:  # wdInlineShapePicture
                image_index += 1
                
                try:
                    # Create a temporary file path
                    temp_image_path = os.path.join(temp_dir, f"temp_img_{image_index}.png")
                    
                    # Select image and copy to clipboard
                    inline_shape.Select()
                    word_app.Selection.Copy()
                    
                    # Use PIL to read image from clipboard
                    try:
                        from PIL import ImageGrab
                        clipboard_image = ImageGrab.grabclipboard()
                        if clipboard_image:
                            # Save to temporary file
                            clipboard_image.save(temp_image_path, format='PNG')
                            # Read image data
                            with open(temp_image_path, 'rb') as f:
                                img_data = f.read()
                            images.append((image_index, img_data, 'png'))
                            # Delete temporary file
                            os.remove(temp_image_path)
                        else:
                            print(f"Warning: Unable to read image {image_index} from clipboard")
                    except Exception as clip_e:
                        print(f"Warning: Error reading image {image_index} from clipboard: {clip_e}")
                except Exception as e:
                    print(f"Warning: Error extracting image {image_index}: {e}")
        
        # If InlineShapes method didn't find images, try finding Shape objects
        if len(images) == 0:
            for shape_idx in range(1, doc.Shapes.Count + 1):
                shape = doc.Shapes(shape_idx)
                # Type = 11 means msoPicture, Type = 13 means msoLinkedPicture
                if shape.Type in [11, 13]:
                    image_index += 1
                    try:
                        temp_image_path = os.path.join(temp_dir, f"temp_img_{image_index}.png")
                        # Select Shape object and copy to clipboard
                        shape.Select()
                        word_app.Selection.Copy()
                        
                        # Use PIL to read image from clipboard
                        try:
                            from PIL import ImageGrab
                            clipboard_image = ImageGrab.grabclipboard()
                            if clipboard_image:
                                clipboard_image.save(temp_image_path, format='PNG')
                                with open(temp_image_path, 'rb') as f:
                                    img_data = f.read()
                                images.append((image_index, img_data, 'png'))
                                os.remove(temp_image_path)
                            else:
                                print(f"Warning: Unable to read image {image_index} from clipboard")
                        except Exception as clip_e:
                            print(f"Warning: Error reading image {image_index} from clipboard: {clip_e}")
                    except Exception as e:
                        print(f"Warning: Error extracting image {image_index}: {e}")
        
        # Clean up temporary directory
        try:
            if os.path.exists(temp_dir) and not os.listdir(temp_dir):
                os.rmdir(temp_dir)
        except:
            pass
        
        doc.Close(SaveChanges=False)
        return doc_text, images
        
    except Exception as e:
        raise Exception(f"Error reading Word document: {e}")
    finally:
        # Close Word application
        word_app.Quit()

def parse_questions(question_section):
    """
    Parse question section, identify question groups and image references
    
    Question group rules:
    - Q1 to Q4 belong to the same question group, presented in order
    - When encountering Q1, always start a new question group
    - When encountering Q2, Q3, Q4:
      * If there is no current group, create a new group (handle missing questions, e.g., only Q3, Q4)
      * If the current group already has a larger question number (e.g., current group has Q4, then encounters Q2), create a new group
      * Otherwise, add to the current group
    """
    lines = question_section.split('\n')
    question_groups = []
    current_group = None
    current_question = None
    group_num = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Match Q1:, Q2:, Q3:, Q4: etc. (support both Chinese and English colons)
        q_match = re.match(r'^Q([1-4])[:：]\s*$', line)
        if q_match:
            question_num = int(q_match.group(1))
            
            # If encountering Q1, always start a new question group
            if question_num == 1:
                group_num += 1
                current_group = {
                    'group': group_num,
                    'questions': {}
                }
                question_groups.append(current_group)
            else:
                # Handle Q2, Q3, Q4 cases
                need_new_group = False
                
                # If there is no current group, need to create a new group
                if current_group is None:
                    need_new_group = True
                else:
                    # If the current group already has a larger question number, need to create a new group
                    existing_questions = current_group['questions'].keys()
                    if existing_questions:
                        max_existing = max(existing_questions)
                        if question_num < max_existing:
                            need_new_group = True
                
                if need_new_group:
                    group_num += 1
                    current_group = {
                        'group': group_num,
                        'questions': {}
                    }
                    question_groups.append(current_group)
            
            # Add question to current group
            current_group['questions'][question_num] = {
                'images': []
            }
            current_question = question_num
        
        # Match image references ![][image1] or ![][imageN]
        img_match = re.match(r'^!\[\]\[image(\d+)\]', line)
        if img_match:
            image_ref = int(img_match.group(1))
            if current_group and current_question:
                current_group['questions'][current_question]['images'].append(image_ref)
    
    return question_groups

def save_image(image_data, image_format, output_path):
    """
    Save image binary data to file
    
    Parameters:
    - image_data: Image binary data
    - image_format: Image format (e.g., 'png', 'jpg')
    - output_path: Output file path
    """
    try:
        # Ensure output directory exists
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Write raw data directly in binary mode
        with open(output_path, 'wb') as f:
            f.write(image_data)
        
        # Verify file was written successfully
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise IOError(f"File write failed: {output_path}")
        
        return output_path
        
    except IOError as e:
        raise IOError(f"File write failed: {e}")
    except Exception as e:
        raise Exception(f"Error occurred while saving image: {e}")

def extract_text_with_ocr(image_path):
    """Extract text from image using OCR, supports Chinese and English"""
    try:
        # Open image
        image = Image.open(image_path)
        
        # Image preprocessing: convert to RGB mode if not already
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Try using Chinese and English OCR (Traditional Chinese + English)
        try:
            text = pytesseract.image_to_string(image, lang='chi_tra+eng', config='--psm 6')
        except Exception:
            # If multilingual fails, try Simplified Chinese + English
            try:
                text = pytesseract.image_to_string(image, lang='chi_sim+eng', config='--psm 6')
            except:
                # Finally try English only
                try:
                    text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')
                except:
                    # If all languages fail, use default language
                    text = pytesseract.image_to_string(image, config='--psm 6')
        
        return text.strip()
    except ImportError:
        print("Warning: pytesseract is not installed, skipping OCR extraction. Please run: pip install pytesseract")
        print("Note: You also need to install Tesseract OCR engine: https://github.com/UB-Mannheim/tesseract/wiki")
        return ""
    except Exception as e:
        error_msg = str(e)
        if "tesseract" in error_msg.lower() and ("not installed" in error_msg.lower() or "not found" in error_msg.lower()):
            print(f"OCR extraction failed for {image_path}: Tesseract OCR is not properly configured.")
            print("Please install Tesseract OCR from: https://github.com/UB-Mannheim/tesseract/wiki")
            print("Or set the path manually: pytesseract.pytesseract.tesseract_cmd = 'path/to/tesseract.exe'")
        else:
            print(f"OCR extraction failed for {image_path}: {e}")
        return ""

def main():
    # Create img folder
    img_dir = 'img'
    if not os.path.exists(img_dir):
        os.makedirs(img_dir)
    
    # Extract images and text from Word document
    # Try .docx format first
    word_file = 'oa.docx'
    if not os.path.exists(word_file):
        # Try .doc format
        word_file = 'oa.doc'
        if not os.path.exists(word_file):
            raise FileNotFoundError("File not found: oa.docx or oa.doc")
    
    print(f"Reading Word document: {word_file}...")
    
    # Choose processing method based on file extension
    if word_file.lower().endswith('.docx'):
        doc_text, extracted_images = extract_images_from_docx(word_file)
    else:
        doc_text, extracted_images = extract_images_from_doc(word_file)
    print(f"Found {len(extracted_images)} images")
    
    # Convert extracted images to dictionary format, using image index as key
    images = {}
    for image_index, image_data, image_format in extracted_images:
        images[image_index] = {
            'format': image_format,
            'data': image_data
        }
    
    # Parse question structure (from Word document text)
    print("Parsing question structure...")
    question_groups = parse_questions(doc_text)
    print(f"Found {len(question_groups)} question groups")
    # Debug info: print question structure
    for group_info in question_groups:
        print(f"  Group {group_info['group']}: {list(group_info['questions'].keys())}")
        for q_num, q_info in group_info['questions'].items():
            print(f"    Q{q_num}: {len(q_info['images'])} images, references: {q_info['images']}")
    
    # Prepare CSV data and collect all OCR text
    csv_data = []
    all_ocr_texts = []
    
    # Process each question group
    for group_info in question_groups:
        group_num = group_info['group']
        questions = group_info['questions']
        
        for question_num in sorted(questions.keys()):
            image_refs = questions[question_num]['images']
            
            for img_idx, image_ref in enumerate(image_refs, 1):
                if image_ref not in images:
                    print(f"Warning: Image image{image_ref} not found")
                    continue
                
                # Generate filename: group_question_image (format: group_question_image.extension)
                filename = f"{group_num}_{question_num}_{img_idx}.{images[image_ref]['format']}"
                image_path = os.path.join(img_dir, filename)
                
                # Save image
                print(f"Saving image: {filename}")
                save_image(
                    images[image_ref]['data'],
                    images[image_ref]['format'],
                    image_path
                )
                
                # Extract text with OCR
                print(f"Extracting text with OCR: {filename}")
                ocr_text = extract_text_with_ocr(image_path)
                
                # Collect OCR text
                all_ocr_texts.append({
                    'filename': filename,
                    'group': group_num,
                    'question': question_num,
                    'image': img_idx,
                    'text': ocr_text
                })
                
                # Add to CSV data (optional, for statistics)
                csv_data.append({
                    'Group': group_num,
                    'Question': question_num,
                    'Image': img_idx,
                    'Image Path': image_path,
                    'OCR Text': ocr_text
                })
    
    # If question parsing didn't find image references, use extracted images in order
    if len(csv_data) == 0 and len(images) > 0:
        print("Warning: Cannot parse question structure from document text, processing images in order...")
        # Assume all images belong to the same group, or assign sequentially to groups
        group_num = 1
        question_num = 1
        img_idx = 0
        
        for image_index in sorted(images.keys()):
            img_idx += 1  # Image number starts from 1
            # Generate filename: group_question_image (format: group_question_image.extension)
            filename = f"{group_num}_{question_num}_{img_idx}.{images[image_index]['format']}"
            image_path = os.path.join(img_dir, filename)
            
            print(f"Saving image: {filename}")
            save_image(
                images[image_index]['data'],
                images[image_index]['format'],
                image_path
            )
            
            print(f"Extracting text with OCR: {filename}")
            ocr_text = extract_text_with_ocr(image_path)
            
            # Collect OCR text
            all_ocr_texts.append({
                'filename': filename,
                'group': group_num,
                'question': question_num,
                'image': img_idx,
                'text': ocr_text
            })
            
            csv_data.append({
                'Group': group_num,
                'Question': question_num,
                'Image': img_idx,
                'Image Path': image_path,
                'OCR Text': ocr_text
            })
    
    # Save all OCR text to a single text file
    output_txt_file = 'ocr_all_results.txt'
    print(f"\nSaving all OCR text to {output_txt_file}...")
    
    try:
        with open(output_txt_file, 'w', encoding='utf-8') as txt_file:
            for idx, item in enumerate(all_ocr_texts, 1):
                # Write separator and image information
                txt_file.write("=" * 80 + "\n")
                txt_file.write(f"Image {idx}: {item['filename']}\n")
                txt_file.write(f"Group: {item['group']}, Question: {item['question']}, Image: {item['image']}\n")
                txt_file.write("-" * 80 + "\n")
                # Write OCR text
                txt_file.write(item['text'])
                txt_file.write("\n\n")
        
        print(f"All OCR text saved to {output_txt_file}")
    except Exception as e:
        print(f"Error: Failed to save text file {output_txt_file}: {e}")
    
    # Save CSV file
    print("Saving CSV file...")
    csv_filename = 'questions_data.csv'
    with open(csv_filename, 'w', newline='', encoding='utf-8-sig') as csvfile:
        fieldnames = ['Group', 'Question', 'Image', 'Image Path', 'OCR Text']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(csv_data)
    
    print(f"Done! Processed {len(csv_data)} records, CSV saved to {csv_filename}")

if __name__ == '__main__':
    main()

